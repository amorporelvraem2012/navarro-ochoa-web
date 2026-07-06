# app_web.py
"""
Navarro & Ochoa Abogados — Formateador de Escritos y Documentos
------------------------------------------------------------------
streamlit run app_web.py

Dependencias Python:
    pip install streamlit python-docx qrcode pillow pytesseract pdf2image

Dependencias de sistema (fuera de Python, necesarias para el OCR):
    - Tesseract OCR      (https://github.com/tesseract-ocr/tesseract)
      Windows: instalador oficial. Linux: sudo apt install tesseract-ocr tesseract-ocr-spa
    - Poppler            (necesario para leer PDFs con pdf2image)
      Windows: https://github.com/oschwartz10612/poppler-windows  Linux: sudo apt install poppler-utils

Alcance de esta herramienta:
    Este sistema NO redacta argumentos legales ni analiza casos. Es un
    FORMATEADOR: toma el texto que tú escribes (hechos, petitorio,
    fundamentos de derecho) y lo organiza en un documento Word con el
    formato, numeración y márgenes de un escrito judicial. El contenido
    jurídico de fondo debe ser redactado por un abogado antes de
    ingresarlo aquí.

Nota legal importante:
    El código QR que se agrega a cada escrito es un CONTROL INTERNO DE
    VERIFICACIÓN del estudio (permite comprobar que un documento salió de
    este sistema y no ha sido alterado). NO es una firma digital
    certificada conforme a la Ley N.º 27269 (Ley de Firmas y Certificados
    Digitales) ni sustituye la firma manuscrita u otra firma electrónica
    reconocida legalmente. Debe visarse/firmarse el documento por los
    medios oficiales correspondientes antes de presentarlo ante una
    instancia judicial, fiscal o administrativa.
"""

import hashlib
import io
import json
from datetime import datetime
from pathlib import Path

import streamlit as st
import streamlit.components.v1 as components
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt
from PIL import Image

import qrcode

# OCR: dependen de binarios de sistema (Tesseract / Poppler).
# Si no están instalados, la app sigue funcionando; solo se avisa al
# entrar a "Procesar OCR".
try:
    import pytesseract
    from pdf2image import convert_from_bytes

    OCR_DISPONIBLE = True
except ImportError:
    OCR_DISPONIBLE = False

# ---------------------------------------------------------------------------
# Configuración general y almacenamiento local
# ---------------------------------------------------------------------------
st.set_page_config(page_title="Navarro & Ochoa Abogados", page_icon="⚖️", layout="wide")

BASE_DIR = Path(__file__).parent
DOCS_DIR = BASE_DIR / "documentos_generados"
CASOS_PATH = BASE_DIR / "casos.json"
SITIO_HTML_PATH = BASE_DIR / "navarro-ochoa-abogados.html"

# Buscar el archivo HTML aunque tenga número en el nombre
if not SITIO_HTML_PATH.exists():
    import glob
    htmls = sorted(glob.glob(str(BASE_DIR / "navarro-ochoa-abogados*.html")))
    if htmls:
        SITIO_HTML_PATH = Path(htmls[-1])  # toma el más reciente
DOCS_DIR.mkdir(exist_ok=True)

# Clave de acceso al sistema interno. Cámbiala antes de usar en producción,
# o reemplaza esta verificación por tu propio sistema de usuarios.
CLAVE_ACCESO_INTERNO = "noa-2026"

TIPOS_PROCESO = {
    "constitucional_amparo": "PROCESO CONSTITUCIONAL DE AMPARO",
    "constitucional_habeas_corpus": "PROCESO CONSTITUCIONAL DE HÁBEAS CORPUS",
    "contencioso_administrativo": "PROCESO CONTENCIOSO ADMINISTRATIVO",
    "civil_conocimiento": "PROCESO CIVIL DE CONOCIMIENTO",
    "penal_denuncia": "DENUNCIA PENAL",
    "recurso_administrativo": "RECURSO ADMINISTRATIVO",
}

JUEZ_POR_TIPO = {
    "constitucional_amparo": "SEÑOR JUEZ CONSTITUCIONAL DE TURNO",
    "constitucional_habeas_corpus": "SEÑOR JUEZ PENAL DE TURNO (HÁBEAS CORPUS)",
    "contencioso_administrativo": "SEÑOR JUEZ CONTENCIOSO ADMINISTRATIVO",
    "civil_conocimiento": "SEÑOR JUEZ CIVIL",
    "penal_denuncia": "SEÑOR FISCAL PROVINCIAL PENAL DE TURNO",
    "recurso_administrativo": "SEÑOR(A) [AUTORIDAD ADMINISTRATIVA COMPETENTE]",
}

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


def fecha_larga_es(fecha: datetime) -> str:
    """Fecha en formato 'DD de <mes> de AAAA', sin depender del locale del sistema."""
    return f"{fecha.day:02d} de {MESES_ES[fecha.month]} de {fecha.year}"


# ---------------------------------------------------------------------------
# Utilidades de almacenamiento (registro interno de casos)
# ---------------------------------------------------------------------------
def cargar_casos() -> list:
    if not CASOS_PATH.exists():
        return []
    try:
        with open(CASOS_PATH, "r", encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return []


def guardar_caso(registro: dict) -> None:
    casos = cargar_casos()
    casos.append(registro)
    with open(CASOS_PATH, "w", encoding="utf-8") as f:
        json.dump(casos, f, ensure_ascii=False, indent=2)


def buscar_por_codigo(codigo: str):
    codigo = codigo.strip().upper()
    for c in cargar_casos():
        if c.get("codigo", "").upper() == codigo:
            return c
    return None


# ---------------------------------------------------------------------------
# Generación del código de verificación y del QR
# ---------------------------------------------------------------------------
def generar_codigo_verificacion(datos: dict) -> str:
    base = "|".join(
        [
            datos["tipo"],
            datos["demandante"],
            datos["demandado"],
            datos["abogado"],
            datos["colegiatura"],
            datetime.now().isoformat(),
        ]
    )
    digest = hashlib.sha256(base.encode("utf-8")).hexdigest()[:10].upper()
    return f"NOA-{digest}"


def generar_imagen_qr(codigo: str) -> io.BytesIO:
    contenido = (
        f"Navarro & Ochoa Abogados\n"
        f"Control interno de verificación\n"
        f"Código: {codigo}\n"
        f"Este código NO constituye firma digital certificada."
    )
    qr = qrcode.QRCode(box_size=8, border=2)
    qr.add_data(contenido)
    qr.make(fit=True)
    img = qr.make_image(fill_color="black", back_color="white")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Generación del escrito judicial (.docx)
# ---------------------------------------------------------------------------
def generar_escrito_docx(datos: dict, codigo: str, qr_buffer: io.BytesIO) -> io.BytesIO:
    doc = Document()

    # Márgenes estándar de escrito judicial peruano
    seccion = doc.sections[0]
    seccion.left_margin = Cm(3)
    seccion.right_margin = Cm(2.5)
    seccion.top_margin = Cm(2.5)
    seccion.bottom_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)

    # --------- ENCABEZADO CON LOGO ---------
    # Intentar cargar el logo desde archivo local
    logo_path = BASE_DIR / "logo-clean.jpg"
    header_section = seccion.header
    header_para = header_section.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.paragraph_format.space_after = Pt(4)

    try:
        if logo_path.exists():
            run_logo = header_para.add_run()
            run_logo.add_picture(str(logo_path), width=Cm(2.2))
    except Exception:
        pass

    run_nombre = header_para.add_run("  NAVARRO & OCHOA ABOGADOS")
    run_nombre.bold = True
    run_nombre.font.size = Pt(13)
    run_nombre.font.name = "Times New Roman"

    run_datos = header_para.add_run("\nEstudio Jurídico · Lima, Perú · +51 990 632 014")
    run_datos.font.size = Pt(8)
    run_datos.font.name = "Times New Roman"

    # Línea separadora en el encabezado
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = header_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '1a4a7a')
    pBdr.append(bottom)
    pPr.append(pBdr)
    # --------- FIN ENCABEZADO ---------

    encabezado_juez = JUEZ_POR_TIPO.get(
        datos["tipo"], "SEÑOR(A) JUEZ / AUTORIDAD COMPETENTE"
    )
    titulo_proceso = TIPOS_PROCESO.get(
        datos["tipo"], datos["tipo"].replace("_", " ").upper()
    )

    # --------- Encabezado (sumilla) ---------
    sumilla = doc.add_paragraph()
    sumilla.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = sumilla.add_run(f"EXPEDIENTE N.° {codigo}\nSUMILLA: {titulo_proceso}")
    r.bold = True
    r.font.size = Pt(10)

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_juez = p.add_run(encabezado_juez)
    run_juez.bold = True
    run_juez.underline = True

    doc.add_paragraph()

    # --------- Datos de las partes ---------
    tabla = doc.add_table(rows=0, cols=2)
    tabla.style = "Table Grid"
    tabla.autofit = True
    filas = [
        ("Materia", titulo_proceso),
        ("Demandante", datos["demandante"]),
        ("Demandado", datos["demandado"]),
        ("Abogado patrocinante", datos["abogado"]),
        ("Registro CAL/colegiatura", datos["colegiatura"]),
        ("Fecha de elaboración", fecha_larga_es(datetime.now())),
    ]
    for etiqueta, valor in filas:
        fila = tabla.add_row()
        celda_izq = fila.cells[0]
        celda_izq.text = etiqueta
        celda_izq.paragraphs[0].runs[0].bold = True
        celda_izq.paragraphs[0].runs[0].font.size = Pt(11)
        celda_der = fila.cells[1]
        celda_der.text = valor
        celda_der.paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.add_run(
        f"{datos['demandante']}, con el patrocinio del abogado que suscribe, "
        f"me presento ante su despacho y digo:"
    )

    # --------- I. Fundamentos de hecho ---------
    h1 = doc.add_paragraph()
    h1.add_run("I. FUNDAMENTOS DE HECHO").bold = True
    for i, parrafo in enumerate(
        (datos["hechos"] or "[Pendiente de completar]").split("\n"), start=1
    ):
        if not parrafo.strip():
            continue
        pf = doc.add_paragraph()
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pf.add_run(f"{i}. {parrafo.strip()}")

    # --------- II. Fundamentos de derecho ---------
    h2 = doc.add_paragraph()
    h2.add_run("II. FUNDAMENTOS DE DERECHO").bold = True
    texto_derecho = (datos.get("fundamentos_derecho") or "").strip()
    if texto_derecho:
        for i, parrafo in enumerate(texto_derecho.split("\n"), start=1):
            if not parrafo.strip():
                continue
            pd = doc.add_paragraph()
            pd.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            pd.add_run(f"{i}. {parrafo.strip()}")
    else:
        pd = doc.add_paragraph()
        pd.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pd.add_run("[Pendiente de completar — ingresa aquí la base legal aplicable a tu caso]")

    # --------- III. Petitorio ---------
    h3 = doc.add_paragraph()
    h3.add_run("III. PETITORIO").bold = True
    for i, parrafo in enumerate(
        (datos["pretensiones"] or "[Pendiente de completar]").split("\n"), start=1
    ):
        if not parrafo.strip():
            continue
        pp = doc.add_paragraph()
        pp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        pp.add_run(f"{i}. {parrafo.strip()}")

    # --------- IV. Anexos ---------
    h4 = doc.add_paragraph()
    h4.add_run("IV. ANEXOS").bold = True
    pa = doc.add_paragraph()
    pa.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pa.add_run(
        "1-A. Se adjuntan los medios probatorios que sustentan lo expuesto "
        "en el presente escrito."
    )

    # --------- Cierre ---------
    doc.add_paragraph()
    cierre = doc.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cierre.add_run(
        "POR LO EXPUESTO: solicito a su despacho tener por presentado el "
        "presente escrito y proveer conforme a lo peticionado."
    )

    doc.add_paragraph()
    fecha_p = doc.add_paragraph()
    fecha_p.add_run(f"Lima, {fecha_larga_es(datetime.now())}.")

    # --------- Firma ---------
    doc.add_paragraph()
    doc.add_paragraph()
    firma = doc.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma.add_run(
        f"_______________________________\n"
        f"{datos['abogado']}\n"
        f"Abogado — Reg. {datos['colegiatura']}"
    )

    # --------- Pie de verificación con QR ---------
    doc.add_paragraph()
    pie_regla = doc.add_paragraph()
    pie_regla.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie_regla.add_run("· " * 30).font.size = Pt(8)

    qr_parrafo = doc.add_paragraph()
    qr_parrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_run = qr_parrafo.add_run()
    qr_run.add_picture(qr_buffer, width=Inches(1.0))

    nota = doc.add_paragraph()
    nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nota_run = nota.add_run(
        f"Navarro & Ochoa Abogados · Código de verificación interna: {codigo}\n"
        f"Este código permite comprobar el origen del documento en el "
        f"registro del estudio. No constituye firma digital certificada."
    )
    nota_run.font.size = Pt(8)
    nota_run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------------------------------------------------------------------
# OCR
# ---------------------------------------------------------------------------
def extraer_texto_ocr(archivo) -> str:
    nombre = archivo.name.lower()
    contenido = archivo.read()

    textos = []
    try:
        if nombre.endswith(".pdf"):
            paginas = convert_from_bytes(contenido)
            for i, pagina in enumerate(paginas, start=1):
                texto_pagina = pytesseract.image_to_string(pagina, lang="spa")
                textos.append(f"--- Página {i} ---\n{texto_pagina}")
        else:
            imagen = Image.open(io.BytesIO(contenido))
            textos.append(pytesseract.image_to_string(imagen, lang="spa"))
    except pytesseract.TesseractNotFoundError:
        return (
            "ERROR: no se encontró el motor Tesseract OCR instalado en este "
            "equipo. Instálalo (ver comentario al inicio de app_web.py) y "
            "vuelve a intentarlo."
        )
    except Exception as exc:  # noqa: BLE001 - mostramos el error al usuario
        return f"ERROR al procesar el archivo: {exc}"

    return "\n\n".join(textos).strip() or "No se detectó texto en el documento."


# ---------------------------------------------------------------------------
# Interfaz
# ---------------------------------------------------------------------------
vista = st.query_params.get("vista", "publico")

# ============================================================================
# VISTA PÚBLICA — la página web del estudio, tal cual, con el botón
# "Acceso interno" en su propio menú (definido en el .html) apuntando a
# esta misma app con ?vista=interno
# ============================================================================
if vista == "publico":
    st.markdown(
        """
        <style>
            .block-container {padding: 0 !important; max-width: 100% !important;}
            header[data-testid="stHeader"] {display: none;}
            iframe {display: block;}
        </style>
        """,
        unsafe_allow_html=True,
    )

    if not SITIO_HTML_PATH.exists():
        st.error(
            f"No se encontró '{SITIO_HTML_PATH.name}'. Colócalo en la misma "
            "carpeta que app_web.py."
        )
    else:
        html_sitio = SITIO_HTML_PATH.read_text(encoding="utf-8")
        components.html(html_sitio, height=6200, scrolling=True)

# ============================================================================
# VISTA INTERNA — el sistema de redacción judicial, protegido con clave
# ============================================================================
else:
    st.markdown(
        "[← Volver al sitio web público](?vista=publico)",
        unsafe_allow_html=False,
    )

    if "autenticado" not in st.session_state:
        st.session_state.autenticado = False

    if not st.session_state.autenticado:
        st.title("🔐 Acceso interno")
        st.write("Este sistema es de uso exclusivo del personal de Navarro & Ochoa Abogados.")
        clave = st.text_input("Clave de acceso", type="password")
        if st.button("Ingresar"):
            if clave == CLAVE_ACCESO_INTERNO:
                st.session_state.autenticado = True
                st.rerun()
            else:
                st.error("Clave incorrecta.")
        st.stop()

    st.title("🛡️ Navarro & Ochoa Abogados")
    st.subheader("Formateador de Escritos y Documentos")
    st.caption(
        "Herramienta de forma, no de fondo: da estructura, numeración y "
        "verificación a lo que tú redactas. No sustituye el criterio "
        "profesional de un abogado."
    )

    menu = st.sidebar.selectbox(
        "Menú", ["Nuevo Escrito", "Procesar OCR", "Validar Certificado", "Mis Casos"]
    )

    # --------------------------- Nuevo Escrito ---------------------------------
    if menu == "Nuevo Escrito":
        col1, col2 = st.columns(2)
        with col1:
            tipo = st.selectbox(
                "Tipo de proceso",
                options=list(TIPOS_PROCESO.keys()),
                format_func=lambda k: TIPOS_PROCESO[k],
            )
            demandante = st.text_input("Demandante")
            demandado = st.text_input("Demandado")
            abogado = st.text_input("Abogado")
            colegiatura = st.text_input("Colegiatura (Reg. CAL)")

        with col2:
            hechos = st.text_area("Hechos", height=140)
            fundamentos_derecho = st.text_area(
                "Fundamentos de Derecho (redáctalos tú o pégalos aquí)",
                height=100,
                placeholder="Ej.: Artículo 200° de la Constitución Política del Perú; artículo 2° del Código Procesal Constitucional...",
            )
            pretensiones = st.text_area("Petitorio / Pretensiones", height=140)

        st.caption(
            "Esta herramienta solo da formato al documento: la numeración, "
            "los márgenes y el QR de verificación se generan automáticamente, "
            "pero **todo el contenido legal (hechos, fundamentos de derecho y "
            "petitorio) debe ser redactado por ti**. El sistema no analiza el "
            "caso ni sugiere argumentos."
        )

        if st.button("Generar Escrito con QR de verificación"):
            campos_obligatorios = {
                "Demandante": demandante,
                "Demandado": demandado,
                "Abogado": abogado,
                "Colegiatura": colegiatura,
                "Hechos": hechos,
                "Pretensiones": pretensiones,
            }
            faltantes = [nombre for nombre, valor in campos_obligatorios.items() if not valor.strip()]

            if faltantes:
                st.error("Completa los siguientes campos antes de generar el escrito: " + ", ".join(faltantes))
            else:
                datos = {
                    "tipo": tipo,
                    "demandante": demandante,
                    "demandado": demandado,
                    "abogado": abogado,
                    "colegiatura": colegiatura,
                    "hechos": hechos,
                    "fundamentos_derecho": fundamentos_derecho,
                    "pretensiones": pretensiones,
                }
                codigo = generar_codigo_verificacion(datos)
                qr_buffer = generar_imagen_qr(codigo)
                qr_buffer_para_docx = io.BytesIO(qr_buffer.getvalue())

                docx_buffer = generar_escrito_docx(datos, codigo, qr_buffer_para_docx)

                nombre_archivo = f"{codigo}_{demandante.replace(' ', '_')}.docx"
                ruta_guardado = DOCS_DIR / nombre_archivo
                with open(ruta_guardado, "wb") as f:
                    f.write(docx_buffer.getvalue())

                guardar_caso(
                    {
                        "codigo": codigo,
                        "fecha": datetime.now().strftime("%d/%m/%Y %H:%M"),
                        "tipo": TIPOS_PROCESO.get(tipo, tipo),
                        "demandante": demandante,
                        "demandado": demandado,
                        "abogado": abogado,
                        "colegiatura": colegiatura,
                        "archivo": str(ruta_guardado.name),
                    }
                )

                st.success(f"Escrito generado. Código de verificación: {codigo}")

                colA, colB = st.columns([1, 3])
                with colA:
                    st.image(qr_buffer, caption="QR de verificación", width=140)
                with colB:
                    st.download_button(
                        "Descargar Documento (.docx)",
                        data=docx_buffer.getvalue(),
                        file_name=nombre_archivo,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

    # --------------------------- Procesar OCR -----------------------------------
    elif menu == "Procesar OCR":
        if not OCR_DISPONIBLE:
            st.warning(
                "Los paquetes 'pytesseract' y/o 'pdf2image' no están instalados. "
                "Ejecuta: pip install pytesseract pdf2image — y asegúrate de tener "
                "instalado Tesseract OCR y Poppler en el sistema (ver comentarios "
                "al inicio de app_web.py)."
            )
        else:
            archivo = st.file_uploader("Sube PDF o imagen escaneada", type=["pdf", "png", "jpg", "jpeg"])
            if archivo and st.button("Extraer Texto con OCR"):
                with st.spinner("Procesando OCR…"):
                    texto = extraer_texto_ocr(archivo)
                st.text_area("Texto Extraído", texto, height=300)
                st.download_button(
                    "Descargar texto (.txt)",
                    data=texto.encode("utf-8"),
                    file_name=f"ocr_{archivo.name.rsplit('.', 1)[0]}.txt",
                )

    # --------------------------- Validar Certificado ----------------------------
    elif menu == "Validar Certificado":
        st.write(
            "Busca un escrito emitido por este sistema a partir de su código de "
            "verificación (impreso junto al QR en el documento)."
        )
        codigo_busqueda = st.text_input("Código de verificación", placeholder="NOA-XXXXXXXXXX")

        if st.button("Validar"):
            if not codigo_busqueda.strip():
                st.error("Ingresa un código de verificación.")
            else:
                registro = buscar_por_codigo(codigo_busqueda)
                if registro:
                    st.success("Código encontrado en el registro interno del estudio.")
                    st.json(registro)
                    ruta_archivo = DOCS_DIR / registro["archivo"]
                    if ruta_archivo.exists():
                        with open(ruta_archivo, "rb") as f:
                            st.download_button(
                                "Descargar documento original",
                                data=f.read(),
                                file_name=registro["archivo"],
                            )
                else:
                    st.error(
                        "Código no encontrado. Este documento no figura en el "
                        "registro de este estudio, o el código fue mal digitado."
                    )

        st.caption(
            "Esta validación solo confirma que el documento fue emitido por este "
            "sistema interno. No es una verificación de firma digital certificada "
            "ante una entidad de certificación acreditada."
        )

    # --------------------------- Mis Casos ---------------------------------------
    elif menu == "Mis Casos":
        casos = cargar_casos()

        filtro = st.text_input("Buscar por demandante, demandado o código")
        if filtro:
            filtro_lower = filtro.lower()
            casos = [
                c
                for c in casos
                if filtro_lower in c.get("demandante", "").lower()
                or filtro_lower in c.get("demandado", "").lower()
                or filtro_lower in c.get("codigo", "").lower()
            ]

        if not casos:
            st.info("Aún no hay escritos generados que coincidan con la búsqueda.")
        else:
            for c in reversed(casos):
                with st.expander(f"{c['codigo']} · {c['demandante']} vs. {c['demandado']} · {c['fecha']}"):
                    st.write(f"**Tipo de proceso:** {c['tipo']}")
                    st.write(f"**Abogado:** {c['abogado']} (Reg. {c['colegiatura']})")
                    ruta_archivo = DOCS_DIR / c["archivo"]
                    if ruta_archivo.exists():
                        with open(ruta_archivo, "rb") as f:
                            st.download_button(
                                "Descargar de nuevo",
                                data=f.read(),
                                file_name=c["archivo"],
                                key=f"descarga_{c['codigo']}",
                            )
                    else:
                        st.warning("El archivo original ya no está disponible en este equipo.")

    st.sidebar.info("Formateador de Documentos - Navarro & Ochoa Abogados")
    if st.sidebar.button("Cerrar sesión"):
        st.session_state.autenticado = False
        st.rerun()
